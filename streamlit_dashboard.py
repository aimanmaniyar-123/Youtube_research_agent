import streamlit as st
import requests
import json
from io import BytesIO
from docx import Document
   


# ------------------------------------------------------------
# Streamlit Config
# ------------------------------------------------------------
st.set_page_config(
    page_title="YouTube Research Agent — FAST MODE",
    layout="wide",
)

API_URL = " https://youtube-research-agent.onrender.com"


# ------------------------------------------------------------
# Scrollable Text Renderer
# ------------------------------------------------------------
def render_scrollable_block(text: str):
    st.markdown(
        f"""
        <div style="
            background-color:#1c1c1c;
            color:#fff;
            padding:15px;
            border-radius:8px;
            max-height:500px;
            overflow-y:scroll;
            white-space:pre-wrap;
            font-size:15px;">
            {text}
        </div>
        """,
        unsafe_allow_html=True,
    )


# ------------------------------------------------------------
# CLEAN JSON → READABLE TEXT REPORT
# ------------------------------------------------------------
def convert_analysis_to_sections(report: dict):
    """Return structured sections instead of raw text lines."""
    def safe(x):
        if x in [None, "", "N/A", {}, [], "null"]:
            return None
        if isinstance(x, dict):
            return json.dumps(x, indent=2)
        return str(x)

    sections = {}

    # EXECUTIVE SUMMARY
    summary = safe(report.get("executive_summary"))
    if summary:
        sections["Executive Summary"] = summary

    # METRICS
    metrics = report.get("metrics", {})
    metric_lines = []
    for k, v in metrics.items():
        val = safe(v)
        if val:
            metric_lines.append(f"{k.replace('_',' ').title()}: {val}")
    if metric_lines:
        sections["Metrics"] = "\n".join(metric_lines)

    # THEMES
    theme_lines = []
    for t in report.get("themes", []):
        if isinstance(t, dict):
            name = safe(t.get("name"))
            if not name:
                continue
            freq = safe(t.get("frequency")) or "N/A"
            eng = safe(t.get("engagement")) or "N/A"
            theme_lines.append(f"- {name} (Freq: {freq}, Engagement: {eng})")
    if theme_lines:
        sections["Themes"] = "\n".join(theme_lines)

    # INSIGHTS
    insight_lines = []
    for i in report.get("insights", []):
        if isinstance(i, dict):
            txt = safe(i.get("text"))
            if not txt:
                continue
            conf = safe(i.get("confidence")) or "N/A"
            cat = safe(i.get("category")) or "N/A"
            insight_lines.append(f"- {txt} [Confidence: {conf}, Category: {cat}]")
    if insight_lines:
        sections["Insights"] = "\n".join(insight_lines)

    # RECOMMENDATIONS
    rec_lines = []
    for r in report.get("recommendations", []):
        if isinstance(r, dict):
            title = safe(r.get("title"))
            desc = safe(r.get("description"))
            if not title or not desc:
                continue
            pr = safe(r.get("priority")) or "N/A"
            imp = safe(r.get("impact")) or "N/A"
            rec_lines.append(f"- {title}: {desc} [Priority: {pr}, Impact: {imp}]")
    if rec_lines:
        sections["Recommendations"] = "\n".join(rec_lines)

    return sections



# ------------------------------------------------------------
# EXPORT HELPERS
# ------------------------------------------------------------
def export_as_text_file(text: str):
    return BytesIO(text.encode("utf-8"))


from docx.shared import RGBColor

def export_as_docx(sections: dict):
    doc = Document()
    for heading, body in sections.items():
        h = doc.add_heading(heading, level=1)
        h.runs[0].font.color.rgb = RGBColor(30, 80, 200)

        for line in body.split("\n"):
            doc.add_paragraph(line)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio



def export_as_html(sections: dict):
    html = """
    <html><head><meta charset="UTF-8">
    <style>
        body { font-family: Arial; padding: 20px; }
        h2 { color: #1e50c8; }
        pre { background: #f4f4f4; padding: 10px; white-space: pre-wrap; }
    </style></head><body>
    <h1 style="color:#1e50c8">YouTube Analysis Report</h1>
    """

    for heading, body in sections.items():
        html += f"<h2>{heading}</h2>"
        html += f"<pre>{body}</pre>"

    html += "</body></html>"
    return BytesIO(html.encode("utf-8"))





# ------------------------------------------------------------
# Title + Input
# ------------------------------------------------------------
st.title("📊 YouTube Research Agent — FAST MODE UI")
st.write("Enter a YouTube Channel URL or @handle or name to analyze instantly.")

channel_input = st.text_input("Enter YouTube Channel URL / Name / @handle")

if st.button("Analyze"):
    if not channel_input.strip():
        st.error("Please enter something.")
        st.stop()

    with st.spinner("Processing..."):
        try:
            response = requests.post(API_URL, json={"query": channel_input}, timeout=120)
            data = response.json()
        except Exception as e:
            st.error(f"Backend error: {e}")
            st.stop()

    # ------------------------------------------------------------
    # CHANNEL SECTION
    # ------------------------------------------------------------
    st.subheader("📌 Channel Information")
    channel = data.get("channel")

    if channel:
        snippet = channel.get("snippet", {})
        stats = channel.get("statistics", {})
        st.json({
            "Title": snippet.get("title"),
            "Description": snippet.get("description"),
            "Subscribers": stats.get("subscriberCount"),
            "Total Views": stats.get("viewCount"),
            "Total Videos": stats.get("videoCount"),
            "Country": snippet.get("country"),
            "Custom URL": snippet.get("customUrl"),
            "Published At": snippet.get("publishedAt"),
        })
    else:
        st.warning("No channel info available.")

    # ------------------------------------------------------------
    # VIDEOS SECTION
    # ------------------------------------------------------------
    st.subheader("🎬 Latest Uploaded Videos")
    videos = data.get("videos", [])
    if videos:
        st.dataframe(videos, use_container_width=True)
    else:
        st.warning("No video metadata found.")

    # ------------------------------------------------------------
    # ANALYSIS SECTION
    # ------------------------------------------------------------
    st.subheader("🧠 AI Analysis Report")
    analysis = data.get("analysis")

    if analysis:
        sections = convert_analysis_to_sections(analysis)

# Render on UI
        for heading, body in sections.items():
            st.markdown(f"## <span style='color:#1e50c8'>{heading}</span>", unsafe_allow_html=True)
            render_scrollable_block(body)
    else:
          st.warning("No report generated.")
          st.stop()

    # ------------------------------------------------------------
    # EXPORT OPTIONS
    # ------------------------------------------------------------
    st.subheader("📁 Export Report")

    col1, col2, col3, col4 = st.columns(4)

    with col1:
        st.download_button(
        "📄 TXT",
        export_as_text_file("\n\n".join([f"{k}\n{v}" for k,v in sections.items()])),
        "youtube_report.txt",
        "text/plain"
    )

    with col2:
        st.download_button(
        "📝 DOCX",
        export_as_docx(sections),
        "youtube_report.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    with col3:
        st.download_button(
        "🌐 HTML",
        export_as_html(sections),
        "youtube_report.html",
        "text/html"
    )

    

    # ------------------------------------------------------------
    # Semantic Info
    # ------------------------------------------------------------
    st.subheader("🔎 Semantic Retrieval Info")
    st.info(f"Semantic Chunks Used: **{data.get('semantic_used', 0)}**")

    st.success("✔ Fast Mode Analysis Completed!")
